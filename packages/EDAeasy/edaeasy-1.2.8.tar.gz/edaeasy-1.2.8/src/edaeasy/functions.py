import os
from typing import Dict, Tuple, Union

import pandas as pd
import pingouin as pg
from docx import Document
from scipy.stats import chi2_contingency, fisher_exact, mannwhitneyu, shapiro, ttest_ind
from statsmodels.stats.contingency_tables import Table, Table2x2
from statsmodels.stats.proportion import proportion_confint


def resumen_datos(dataframe: pd.DataFrame) -> tuple[pd.DataFrame, tuple[int, int]]:
    """Generate a summary DataFrame of the input DataFrame 'dataframe'.

    Parameters
    ----------
    dataframe : pandas.DataFrame
        The input DataFrame for which the summary needs to be generated.

    Returns
    -------
    pandas.DataFrame
        A DataFrame containing summary information for each column in 'df':
        - Type: Data type of the column.
        - Min: Minimum value in the column.
        - Max: Maximum value in the column.
        - Nan %: Percentage of NaN values in the column.
        - # Unique Values: Total number of unique values in the column.
        - Unique values: List of unique values in the column.

    Tuple(Rows, Columns)

    Example
    -------
    >>> data = {
            'age': ['[40-50)', '[60-70)', '[70-80)'],
            'time_in_hospital': [8, 3, 5],
            'n_lab_procedures': [72, 34, 45],
            ...
        }
    >>> dataframe = pd.DataFrame(data)
    >>> result = dataframe_summary(df)
    >>> print(result)
               Type       Min        Max  Nan %  # Unique Values                                  Unique values
    Variables
    age       object   [40-50)    [90-100)    0.0        3      ['[70-80)', '[50-60)', '[60-70)', '[40-50)', '[80-90)', ...
    time_in_hospital  int64    1           14    0.0        3        [8, 3, 5]
    n_lab_procedures  int64    1          113    0.0        3        [72, 34, 45]
    ...

    Note
    ----
    The function uses vectorized operations to improve performance and memory usage.
    """
    ret = pd.DataFrame(
        columns=["Type", "Min", "Max", "Nan %", "# Unique Values", "Unique values"]
    )

    for col, content in dataframe.items():
        values = []
        dtype = content.dtype

        # Convert 'object' columns to appropriate data types
        if dtype == "object":
            if col in ["Min", "Max", "Type"]:
                # Clean and convert 'Min', 'Max', and 'Type' columns to strings
                content = content.astype(str)

        values.append(content.dtype)  # Data type after conversion

        try:
            values.append(content.min())  # Min
            values.append(content.max())  # Max
        except Exception:
            values.append("None")
            values.append("None")

        values.append(content.isnull().mean() * 100)  # % of NaN's
        # Calculate the number of unique values in the column
        num_unique_values = len(content.drop_duplicates())
        values.append(num_unique_values)  # Number of unique values
        # Handle the 'Unique values' column as a list of strings
        unique_values = content.drop_duplicates().astype(str).tolist()
        unique_values.sort()
        values.append(unique_values)
        ret.loc[col] = values

    ret.index.names = ["Variables"]
    rows = dataframe.shape[0]
    columns = dataframe.shape[1]
    return ret, (rows, columns)


def tabulacion_frecuencias_proporciones(
    data: pd.DataFrame, column_name: str, show_col=False
) -> pd.DataFrame:
    """
    Calcula la tabulación de frecuencias y proporciones con intervalos de confianza para una columna específica de un DataFrame.

    Parameters:
    data (pd.DataFrame): El DataFrame que contiene los datos.
    column_name (str): El nombre de la columna para la cual se calcularán las proporciones.

    Returns:
    pd.DataFrame: DataFrame con las categorías, número de observaciones, proporciones e intervalos de confianza (CI low - CI upp).
    """
    # Tabulación de frecuencias
    tabulated_data = data[column_name].value_counts().reset_index()
    tabulated_data.columns = ["Value", "Frequency"]

    # Crear una lista para almacenar los resultados de proporciones
    results = []
    if show_col:
        results.append([column_name, "", "", "", ""])

    for category in tabulated_data["Value"]:
        # Filtrar el DataFrame para la categoría específica
        subset = data[data[column_name] == category]
        num_success = subset.shape[0]
        total_samples = len(data[column_name])

        # Calcular la proporción y el intervalo de confianza
        prop = num_success / total_samples
        conf_int = proportion_confint(num_success, total_samples, method="normal")
        ci_lower, ci_upper = conf_int

        # Almacenar los resultados
        results.append(
            [category, num_success, prop * 100, ci_lower * 100, ci_upper * 100]
        )

    # Crear un DataFrame a partir de la lista de resultados
    result_df = pd.DataFrame(
        results,
        columns=["Category", "Number of obs", "Proportion", "CI Lower", "CI Upper"],
    )

    return result_df


def resumen_frecuencias_proporciones(df: pd.DataFrame, show_col=False) -> pd.DataFrame:
    """
    Aplica la función tabulacion_frecuencias_proporciones a todas las columnas de un DataFrame y devuelve un único DataFrame con el resumen de todas las columnas.

    Parameters:
    df (pd.DataFrame): El DataFrame que contiene los datos.

    Returns:
    pd.DataFrame: DataFrame con el resumen de frecuencias y proporciones para todas las columnas.
    """
    result_list = []

    for column in df.columns:
        result = tabulacion_frecuencias_proporciones(df, column, show_col)

        result_list.append(result)

    summary_df = pd.concat(result_list).reset_index(drop=True)

    def highlight_rows(row):
        if row["Category"] in df.columns:
            return ["background-color: #cceeff; color: black"] * len(
                row
            )  # Fondo azul claro para las filas con el nombre de la columna
        else:
            return [""] * len(row)  # Sin estilo para las demás filas

    # Aplicar el estilo solo si se incluyó el nombre de la columna
    if show_col:
        return summary_df.style.apply(highlight_rows, axis=1)
    return summary_df


def calculo_prevalencia(
    disease_cases: int, pop_size: int
) -> Tuple[float, Tuple[float, float]]:
    """
    Calcula la prevalencia de una enfermedad y su intervalo de confianza.

    Parameters:
    disease_cases (int): Número de casos de la enfermedad.
    pop_size (int): Tamaño de la población.

    Returns:
    Tuple[float, Tuple[float, float]]: La prevalencia y el intervalo de confianza (inferior, superior).
    """
    prevalencia = (disease_cases / pop_size) * 100
    ci = proportion_confint(
        count=disease_cases, nobs=pop_size, alpha=0.05, method="wilson"
    )
    return prevalencia, (ci[0] * 100, ci[1] * 100)


# Creacion de reportes en documento de Word
def _verificar_o_crear_documento(ruta_documento: str) -> Document:
    """
    Verifica si el documento de Word existe en la ruta dada. Si existe, lo abre; de lo contrario, crea uno nuevo.

    Parameters:
    ruta_documento (str): Ruta del documento de Word.

    Returns:
    Document: Documento de Word abierto o nuevo.
    """
    if os.path.exists(ruta_documento):
        doc = Document(ruta_documento)
    else:
        doc = Document()
    return doc


def _agregar_tabla_a_documento(df: pd.DataFrame, titulo: str, doc: Document) -> None:
    """
    Agrega una tabla a un documento de Word con los datos de un DataFrame.

    Parameters:
    df (pd.DataFrame): DataFrame con los datos a agregar.
    titulo (str): Título para la sección de la tabla.
    doc (Document): Documento de Word al cual se le agregará la tabla.
    """
    # Agrega un encabezado para la sección
    doc.add_heading(f"Caracterización {titulo.capitalize()}", level=2)

    # Crea una tabla en el documento con el número adecuado de filas y columnas
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Agrega los encabezados de columna a la tabla
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]

    # Agrega los datos del DataFrame a la tabla
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            value = df.iloc[i, j]
            table.cell(i + 1, j).text = (
                str(round(value, 4)) if isinstance(value, (int, float)) else str(value)
            )


def generar_documento(
    resultados_finales: Dict[str, pd.DataFrame], ruta_documento: str
) -> None:
    """
    Genera un documento de Word con tablas para cada DataFrame en el diccionario proporcionado.

    Parameters:
    resultados_finales (Dict[str, pd.DataFrame]): Diccionario con títulos como claves y DataFrames como valores.
    ruta_documento (str): Ruta donde se guardará el documento de Word.
    """
    doc = _verificar_o_crear_documento(ruta_documento)

    for titulo, df in resultados_finales.items():
        _agregar_tabla_a_documento(df, titulo, doc)

    doc.save(ruta_documento)


def _compare_numeric_columns(
    dataframe: pd.DataFrame, col: str, group_col: str
) -> Dict[str, Union[str, float, pd.DataFrame]]:
    """
    Compara columnas numéricas entre dos grupos y devuelve estadísticas y resultados de pruebas.

    Parameters:
    dataframe (pd.DataFrame): DataFrame con los datos.
    col (str): Nombre de la columna numérica a comparar.
    group_col (str): Nombre de la columna que define los grupos.

    Returns:
    Dict[str, Union[str, float, pd.DataFrame]]: Resultados de las pruebas estadísticas y tabla descriptiva.
    """
    import pingouin as pg

    group1 = dataframe[dataframe[group_col] == dataframe[group_col].unique()[0]][
        col
    ].dropna()
    group2 = dataframe[dataframe[group_col] == dataframe[group_col].unique()[1]][
        col
    ].dropna()

    try:
        _, p_group1 = shapiro(group1)
    except Exception:
        p_group1 = 0
    try:
        _, p_group2 = shapiro(group2)
    except Exception:
        p_group2 = 0

    normal_group1 = p_group1 > 0.05
    normal_group2 = p_group2 > 0.05

    try:
        if normal_group1 and normal_group2:
            test_name = "T-test"
            stat, p_value = ttest_ind(group1, group2)
            test_pingouin = pg.ttest(group1, group2).iloc[0]
        else:
            test_name = "Mann-Whitney U"
            stat, p_value = mannwhitneyu(group1, group2)
            test_pingouin = pg.mwu(group1, group2).iloc[0]
        p_value = f"{p_value:.2f}*" if p_value < 0.05 else f"{p_value:.2f}"
    except Exception:
        test_name = None
        stat = None
        p_value = None
        test_pingouin = None

    desc_group1 = group1.describe()
    desc_group2 = group2.describe()
    descriptive_table = pd.DataFrame(
        {
            dataframe[group_col].unique()[0]: desc_group1,
            dataframe[group_col].unique()[1]: desc_group2,
        }
    )

    return {
        "col_name": col,
        "test_name": test_name,
        "p_value": p_value,
        "statsmodels_result": {"statistic": stat, "p-value": p_value},
        "pingouin_result": test_pingouin,
        "descriptive_table": descriptive_table,
    }


def _compare_categorical_columns(
    dataframe: pd.DataFrame, col1: str, col2: str
) -> Dict[str, Union[str, float, pd.DataFrame]]:
    """
    Compara columnas categóricas y devuelve estadísticas y resultados de pruebas.

    Parameters:
    dataframe (pd.DataFrame): DataFrame con los datos.
    col1 (str): Nombre de la primera columna categórica.
    col2 (str): Nombre de la segunda columna categórica.

    Returns:
    Dict[str, Union[str, float, pd.DataFrame]]: Resultados de las pruebas estadísticas y tabla descriptiva.
    """

    contingency_table = pd.crosstab(dataframe[col1], dataframe[col2])

    if contingency_table.shape == (1, 2):
        test_name = None
        p_value = None
        result_statsmodels = None
        result_pingouin = None
    elif contingency_table.shape == (2, 2):
        if (contingency_table < 5).any().any():
            test_name = "Fisher's Exact Test"
            oddsratio, p_value = fisher_exact(contingency_table)
            p_value = f"{p_value:.2f}*" if p_value < 0.05 else f"{p_value:.2f}"
            result_statsmodels = {"oddsratio": oddsratio, "p_value": p_value}
            result_pingouin = None
        else:
            test_name = "Chi-Square Test"
            chi2, p_value, dof, expected = chi2_contingency(contingency_table)
            result_statsmodels = Table2x2(
                contingency_table.values
            ).test_nominal_association()
            result_pingouin = pg.chi2_independence(
                dataframe, col1, col2, correction=False
            )[2]
    else:
        if (contingency_table < 5).any().any():
            test_name = "Fisher's Exact Test"
            bunch = Table(contingency_table).test_nominal_association()
            p_value = (
                f"{bunch.pvalue:.2f}*" if bunch.pvalue < 0.05 else f"{bunch.pvalue:.2f}"
            )
            result_statsmodels = {"chi2-stat": bunch.statistic, "p_value": bunch.pvalue}
            result_pingouin = None
        else:
            test_name = "Chi-Square Test"
            chi2, p_value, dof, expected = chi2_contingency(contingency_table)
            result_statsmodels = Table(
                contingency_table.values
            ).test_nominal_association()
            result_pingouin = pg.chi2_independence(
                dataframe, col1, col2, correction=False
            )[2]

    descriptive_table = contingency_table.describe()

    return {
        "col_name": (col1, col2),
        "test_name": test_name,
        "p_value": p_value,
        "statsmodels_result": result_statsmodels,
        "pingouin_result": result_pingouin,
        "descriptive_table": descriptive_table,
    }


def _get_categories(df: pd.DataFrame, agrupador: str, valor: str) -> list:
    """
    Obtiene las categorías y proporciones para una columna específica agrupada por otra columna.

    Parameters:
    df (pd.DataFrame): DataFrame con los datos.
    agrupador (str): Nombre de la columna de agrupación.
    valor (str): Nombre de la columna de valores.

    Returns:
    list: Lista de listas con categorías, observaciones y proporciones.
    """
    assert agrupador in df.columns, f"La columna {agrupador} no existe en el DataFrame."
    assert valor in df.columns, f"La columna {valor} no existe en el DataFrame."
    cats = sorted(df[valor].dropna().unique())
    targets = [[], []]

    grouped = df.groupby([agrupador, valor]).size().unstack(fill_value=0)
    proportions = grouped.div(grouped.sum(axis=1), axis=0)

    for i, target in enumerate(targets):
        group = grouped.iloc[i]
        proportion = proportions.iloc[i]
        for cat in cats:
            obs = group.get(cat, 0)
            prop = round(proportion.get(cat, 0), 3) * 100
            target.append([cat, obs, prop])

    return targets


def _add_row(
    df: pd.DataFrame, variable: str, neg_class: str, pos_class: str, pval: str
) -> None:
    """
    Agrega una fila a un DataFrame.

    Parameters:
    df (pd.DataFrame): DataFrame al que se añadirá la fila.
    variable (str): Nombre de la variable.
    ncy (str): Valor para la columna ncy.
    cy (str): Valor para la columna cy.
    pval (str): Valor para la columna pval.
    """
    df.loc[len(df)] = [variable, neg_class, pos_class, pval]


def _extract_val(
    r_dict: Dict,
    neg_class: int,
    pos_class: int,
    test_name: str,
    subset: int = 1,
    total_samples: int = 1,
) -> Tuple[str, str, str]:
    """
    Extrae valores de un diccionario de resultados y los formatea.

    Parameters:
    r_dict (Dict): Diccionario con resultados.
    cat_neg (int): Categoría negativa.
    cat_pos (int): Categoría positiva.
    test_name (str): Nombre de la prueba.
    subset (int, optional): Subconjunto de datos. Por defecto es 1.
    total_samples (int, optional): Total de muestras. Por defecto es 1.

    Returns:
    Tuple[str, str, str]: Valores formateados.
    """
    vals = {neg_class: None, pos_class: None}

    for val in vals.keys():
        if test_name == "Chi-Square Test":
            medida = r_dict["descriptive_table"][val]["count"]
            proportion = subset / total_samples
            nq1 = r_dict["descriptive_table"][val]["25%"]
            nq3 = r_dict["descriptive_table"][val]["75%"]
            text_norm = f"{medida} ({proportion:.2f}%) [{nq1:.2f} - {nq3:.2f}]"
        elif r_dict["test_name"] == test_name:
            medida = r_dict["descriptive_table"][val]["mean"]
            nq1 = r_dict["descriptive_table"][val]["25%"]
            nq3 = r_dict["descriptive_table"][val]["75%"]
            text_norm = f"{medida:.2f} [{nq1:.2f} - {nq3:.2f}]"
        else:
            medida = r_dict["descriptive_table"][val]["50%"]
            nq1 = r_dict["descriptive_table"][val]["min"]
            nq3 = r_dict["descriptive_table"][val]["max"]
            text_norm = f"{medida:.2f} [{nq1:.2f} - {nq3:.2f}]"
        pval = r_dict["p_value"] if r_dict["p_value"] is not None else "NA"
        vals[val] = text_norm

    return vals[neg_class], vals[pos_class], pval


def _handle_numeric_column(
    df: pd.DataFrame, col: str, testing_col: str, target_df: pd.DataFrame
) -> None:
    """
    Maneja la comparación de columnas numéricas y agrega los resultados al DataFrame objetivo.

    Parameters:
    df (pd.DataFrame): DataFrame con los datos originales.
    col (str): Nombre de la columna numérica.
    testing_col (str): Nombre de la columna de prueba.
    target_df (pd.DataFrame): DataFrame objetivo donde se resumirán los datos.
    """
    result = _compare_numeric_columns(df, col, testing_col)
    ncy, cy, pval = _extract_val(result, 0, 1, "T-test")
    _add_row(target_df, col, ncy, cy, pval)


def _compute_tabulate(df, col):
    counts = df[col].value_counts().reset_index()
    counts.columns = ["Category", "Number of obs"]
    total = counts["Number of obs"].sum()
    counts["Proportion"] = counts["Number of obs"] / total * 100
    return counts


def _handle_categorical_column(
    df: pd.DataFrame, col: str, testing_col: str, target_df: pd.DataFrame
) -> None:
    """
    Maneja la comparación de columnas categóricas y agrega los resultados al DataFrame objetivo.

    Parameters:
    df (pd.DataFrame): DataFrame con los datos originales.
    col (str): Nombre de la columna categórica.
    testing_col (str): Nombre de la columna de prueba.
    target_df (pd.DataFrame): DataFrame objetivo donde se resumirán los datos.
    """
    test_name = "Chi-Square Test"
    totalcat = df[col].nunique()
    result = _compare_categorical_columns(df, col, testing_col)
    ncy, cy, pval = _extract_val(result, 0, 1, test_name)
    if totalcat > 2:
        _add_row(target_df, col, "", "", pval)
        negativo, positivo = _get_categories(df, testing_col, col)
        for group in zip(negativo, positivo):
            negativo = f"{group[0][1]} ({group[0][2]:.1f})"
            positivo = f"{group[1][1]} ({group[1][2]:.1f})"
            _add_row(target_df, group[0][0], negativo, positivo, "")
    else:
        result_ncy = _compute_tabulate(df.query(f"`{testing_col}` == 0"), col)
        result_cy = _compute_tabulate(df.query(f"`{testing_col}` == 1"), col)
        ncy = "0 (0.0)"
        cy = "0 (0.0)"
        for _, row in result_ncy.iterrows():
            if row["Category"] == 1:
                ncy = f"{row['Number of obs']} ({row['Proportion']:.1f})"
        for _, row in result_cy.iterrows():
            if row["Category"] == 1:
                cy = f"{row['Number of obs']} ({row['Proportion']:.1f})"
        _add_row(target_df, col, ncy, cy, pval)


def paired_descriptive_analysis(
    df: pd.DataFrame, target_df: pd.DataFrame, testing_col: str
) -> None:
    """
    Resume las columnas de un DataFrame en un DataFrame objetivo utilizando pruebas estadísticas.

    Parameters:
    df (pd.DataFrame): DataFrame con los datos originales.
    target_df (pd.DataFrame): DataFrame objetivo donde se resumirán los datos.
    testing_col (str): Nombre de la columna de prueba.
    """
    for col in df.columns:
        if col in df.select_dtypes(include=["number"]).columns:
            _handle_numeric_column(df, col, testing_col, target_df)
        elif col in df.select_dtypes(include=["object", "category"]).columns:
            _handle_categorical_column(df, col, testing_col, target_df)

