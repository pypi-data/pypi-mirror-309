import fitz
from docx import Document
from docx.shared import Inches
import os

SUPPORTED_FORMATS = {
    "input": [".pdf"],
    "output": [".docx"]
}


def convert(input_file, output_file):
    """Convert a PDF file to DOCX."""
    if not input_file.endswith(".pdf") or not output_file.endswith(".docx"):
        raise ValueError("Input must be a .pdf file and output must be a .docx file")

    doc = Document()

    pdf_document = fitz.open(input_file)

    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]

        # Extract text
        text = page.get_text("text")
        if text.strip():
            doc.add_paragraph(text)

        # Extract images
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_filename = f"temp_page{page_num+1}_img{img_index+1}.{image_ext}"

            # Save image temporarily
            with open(image_filename, "wb") as img_file:
                img_file.write(image_bytes)

            # Add the image to the DOCX file
            doc.add_picture(image_filename, width=Inches(5))
            os.remove(image_filename)  # Clean up the temporary image

    # Save the DOCX file
    doc.save(output_file)
