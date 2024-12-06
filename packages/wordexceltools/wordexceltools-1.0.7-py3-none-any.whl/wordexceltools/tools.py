# coding=utf-8
"""
@File    : tools.py
@Time    : 2024/2/29 11:21
@Author  : MengYue Sun ldspdvsun@gmail.com
@Description : 基于docx、python-docx和openpyxl的word文档工具类
"""

import os
import openpyxl
import string
import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx import Document
# 在doc开头调用进行设置
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT


def excel_copy_file(source_excel, target_excel):
    """
    复制 Excel 文件到目标路径。如果目标文件已存在，提示并终止复制。
    :param source_excel: 源文件路径。
    :param target_excel: 目标文件路径。
    :return: True 表示复制成功，False 表示失败。
    """
    try:
        # 检查源文件是否存在
        if not os.path.exists(source_excel):
            raise FileNotFoundError(f"Source file '{source_excel}' does not exist.")

        # 检查目标文件是否已存在
        if os.path.exists(target_excel):
            print(f"Target file '{target_excel}' already exists. Aborting copy operation.")
            return False

        # 加载源文件内容
        workbook = openpyxl.load_workbook(source_excel)

        # 保存到目标文件
        workbook.save(target_excel)
        print(f"Excel file copied to: {target_excel}")
        return True
    except Exception as e:
        print(f"An error occurred in copy_file: {e}")
        return False


def excel_delete_file(target_excel):
    """
    删除指定的 Excel 文件。
    :param target_excel: 目标文件路径。
    :return: True 表示删除成功，False 表示失败。
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(target_excel):
            raise FileNotFoundError(f"File '{target_excel}' does not exist.")

        # 删除文件
        os.remove(target_excel)
        print(f"Excel file deleted: {target_excel}")
        return True
    except Exception as e:
        print(f"An error occurred in delete_file: {e}")
        return False



def word_docxinitial(doc, text_font_type='宋体', text_font_size=12, text_font_line_spacing=30,
                     text_header='Header', text_footer='Footer'):
    """
    设置文档的全局字体样式、大小和页眉页脚。

    :param doc: Document对象，要设置样式的Word文档对象。
    :param text_font_type: str，可选，正文的字体类型，默认为'宋体'。
    :param text_font_size: int，可选，正文的字体大小，默认为12磅。
    :param text_font_line_spacing: int，可选，正文的行间距，默认为30磅。
    :param text_header: str，可选，页眉内容，默认为'Header'。
    :param text_footer: str，可选，页脚内容，默认为'Footer'。

    :return: Document对象，设置完样式后的Word文档对象。
    """
    # 设置正文字体类型、大小和行间距
    doc.styles["Normal"].font.name = text_font_type
    doc.styles["Normal"].font.size = Pt(text_font_size)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    doc.styles["Normal"]._element.line_spacing = Pt(text_font_line_spacing)  # 设置行间距为1.5倍

    # 设置页眉
    header = doc.sections[0].header
    pheader = header.paragraphs[0]  # 获取页眉的第一个段落
    ph_header = pheader.add_run(text_header)
    ph_header.font.name = text_font_type  # 设置页眉字体样式
    ph_header._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    ph_header.font.size = Pt(text_font_size)  # 设置页眉字体大小
    pheader.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 设置页脚
    footer = doc.sections[0].footer
    pfooter = footer.paragraphs[0]  # 获取页脚的第一个段落
    ph_footer = pfooter.add_run(text_footer)
    ph_footer.font.name = text_font_type  # 设置页脚字体样式
    ph_footer._element.rPr.rFonts.set(qn('w:eastAsia'), text_font_type)
    ph_footer.font.size = Pt(text_font_size)  # 设置页脚字体大小
    pfooter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 返回文档对象
    return doc


def word_setsectionformat(doc):
    """
    在doc结尾处调用进行设置
    通过sections（节）进行设置时，节对应文档中的每一页,每个节在没输入内容之前是不存在的,因此在最后才对每个节逐一进行设置
    :param doc:
    :return:
    """
    for sec in doc.sections:
        # 设置页面边距(左上25毫米，右下15毫米)
        sec.top_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        # 设置纸张大小(A4)
        sec.page_height = Mm(297)
        sec.page_width = Mm(210)
        # 设置页眉页脚距离
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(0.2)


def word_add_figure_caption(document, caption_text):
    """
    在 Word 文档中添加图序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith('Heading'):
                section_text = get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split('.')[0]
                break

        # 创建一个新的段落对象用于存放图序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的图序号题注对象
        caption = OxmlElement('w:fldSimple')
        caption.set(qn('w:instr'), f'SEQ Figure \\* ARABIC')
        run = p.add_run()
        run.text = f"图 {section_number}-{table_count + 1}：{caption_text}"
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_figure_caption: {e}")


def word_add_custom_heading(document, level, text):
    """
    向文档添加自定义级别的标题。
    :param document: 文档对象。
    :param level: 标题级别，整数值，1 表示一级标题，2 表示二级标题，以此类推。
    :param text: 标题文本。
    """
    try:
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        else:
            heading_style = f"Heading {level}"
            document.add_heading(text, level=level).style = heading_style

    except ValueError as ve:
        print(f"ValueError in add_custom_heading: {ve}")


def word_add_custom_paragraph(document, text, style=None):
    """
    向文档中添加自定义样式的段落。
    :param document: Document对象，表示要添加段落的文档。
    :param text: 要添加的文本内容。
    :param style: 要应用的样式名称，默认为None。
    """
    try:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)

    except Exception as e:
        print(f"An error occurred in add_custom_paragraph: {e}")


def word_add_custom_table(document, row_count, column_count, style="Table Grid", font_size=10.5, spacing_before=0):
    """
    向文档中添加自定义样式的表格。
    :param document: Document对象，表示要添加表格的文档。
    :param row_count: 表格的行数。
    :param column_count: 表格的列数。
    :param style: 要应用的表格样式名称，默认为 "Table Grid"。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    """
    try:
        table = document.add_table(rows=row_count, cols=column_count)
        if style:
            table.style = style
        else:
            table.style = "Table Grid"
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(spacing_before)  # 设置段前间距
        return table

    except Exception as e:
        print(f"An error occurred in add_custom_table: {e}")


def word_add_indented_paragraph(document, text):
    """
    在 Word 文档中添加一个带有序号和缩进的段落。
    :param document: Word 文档对象
    :param text: 段落文本
    """
    try:
        # 将文本拆分成行
        lines = text.split("\n")

        # 添加段落，并在每一行前添加序号和制表符缩进
        for i, line in enumerate(lines, start=1):
            indented_line = f"\t{i}.{line}"
            document.add_paragraph(indented_line, style="BodyText")

    except Exception as e:
        print(f"An error occurred in add_indented_paragraph: {e}")


def word_add_ordered_list(document, items, index_style='number', indentation=1):
    """
    向文档中添加有序列表，并设置缩进。
    :param document: 文档对象。
    :param items: 有序列表项的列表。
    :param index_style: 索引样式，可以是 'number'（数字）、'lower_alpha'（小写字母）、'upper_alpha'（大写字母）、'roman'（罗马数字）之一，默认为 'number'。
    :param indentation: 缩进量，默认为1级0.5英寸。
    """
    try:
        indexes = []  # 存储索引列表
        if index_style == 'number':
            indexes = [str(i) for i in range(1, len(items) + 1)]
        elif index_style == 'lower_alpha':
            indexes = list(string.ascii_lowercase)[:len(items)]  # 将数字转换为小写字母
        elif index_style == 'upper_alpha':
            indexes = [string.ascii_lowercase[i].upper() for i in range(len(items))]  # 将数字转换为大写字母
        elif index_style == 'roman':
            indexes = [set_number_to_roman(i) for i in range(1, len(items) + 1)]  # 使用 to_roman 方法转换数字为罗马数字
        else:
            raise ValueError("Invalid index style. Use 'number', 'low_alpha', 'upper_alpha', or 'roman'.")

        for index, item in zip(indexes, items):
            # 如果不是最后一项，在末尾添加分号
            if index != indexes[-1]:
                paragraph = document.add_paragraph(f"{index}、{item}；")
            else:
                # 如果是最后一项，在末尾添加句号
                paragraph = document.add_paragraph(f"{index}、{item}。")
            paragraph.paragraph_format.left_indent = Inches(indentation * 0.5)

    except Exception as e:
        print(f"An error occurred in add_ordered_list: {e}")


def word_add_table_caption(document, caption_text):
    """
    在 Word 文档中添加表序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith('Heading'):
                section_text = get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split('.')[0]
                break

        # 创建一个新的段落对象用于存放表序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的表序号题注对象
        caption = OxmlElement('w:fldSimple')
        caption.set(qn('w:instr'), f'SEQ Table \\* ARABIC')
        run = p.add_run()
        run.text = f"\n表 {section_number}-{table_count + 1} {caption_text}"
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_table_caption: {e}")


def excel_get_sheet_cell_value(sheet, row, column):
    """
    获取指定单元格的值，并将其转换为字符串返回。
    参数：
    sheet: 工作表对象
    row: 行索引
    column: 列索引

    返回值：
    单元格的值的字符串表示形式
    """
    try:
        return str(sheet.cell(row, column).value)
    except Exception as e:
        print(f"An error occurred in get_sheet_cell_value: {e}")
        return ""


def excel_get_sheet_cell_date_value(sheet, row, column):
    """
    获取指定单元格的日期值，并将其格式化为字符串返回。
    参数：
    sheet: 工作表对象
    row: 行索引
    column: 列索引
    返回值：
    单元格的日期值的字符串表示形式（格式为'%Y年%m月%d日'）
    """
    try:
        date_value = sheet.cell(row, column).value
        if isinstance(date_value, datetime.datetime):
            return date_value.strftime('%Y年%m月%d日')
        else:
            return str(date_value)
    except Exception as e:
        print(f"An error occurred in get_sheet_cell_date_value: {e}")
        return ""


def excel_get_column_keyword_count(sheet, column_name, keyword):
    """
    遍历指定工作簿中指定工作表的指定列，并返回包含指定关键词的单元格数量。
    :param sheet_name: 工作表对象
    :param column_name: 列标题名称
    :param keyword: 要统计的关键词
    :return: 选择列中包含指定关键词的单元格数量
    """

    try:
        # 获取列标题所在行
        header_row = sheet[1]

        # 根据列标题名称确定列索引
        column_index = None
        for cell in header_row:
            if cell.value == column_name:
                column_index = cell.column
                break

        # 如果未找到列标题，则抛出异常
        if column_index is None:
            raise ValueError(f"Column '{column_name}' not found in sheet '{sheet_name}'")

        # 统计包含指定关键词的单元格数量
        keyword_count = sum(
            1 for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=column_index, max_col=column_index)
            if
            row[0].value is not None and str(row[0].value) == keyword)

        return keyword_count

    except Exception as e:
        print(f"An error occurred in get_column_keyword_count: {e}")
        return -999  # 返回0表示发生了错误，无法统计关键词数量


def excel_get_column_len(sheet, column_name):
    """
    遍历指定工作簿中指定工作表的指定列，并返回该列的长度。
    :param sheet: 工作表对象
    :param column_name: 列标题名称
    :return: 选择列的长度
    """

    try:
        # 获取列标题所在行
        header_row = sheet[1]

        # 根据列标题名称确定列索引
        column_index = None
        for cell in header_row:
            if cell.value == column_name:
                column_index = cell.column
                break

        # 如果未找到列标题，则抛出异常
        if column_index is None:
            raise ValueError(f"Column '{column_name}' not found in sheet '{sheet_name}'")

        # 使用集合来跟踪已经出现过的值
        unique_values = set()

        # 遍历指定列的所有数据，从第二行开始
        column_length = 0
        for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=column_index, max_col=column_index):
            value = row[0].value
            # 如果值不为空且未出现过，则增加计数器并将值添加到集合中
            if value is not None and value not in unique_values:
                column_length += 1
                unique_values.add(value)

        return column_length

    except Exception as e:
        print(f"An error occurred in get_column_len: {e}")
        return -999  # 返回0表示发生了错误，无法统计列长度


def excel_get_column_list(sheet, column_name):
    """
    遍历指定工作簿中指定工作表的指定列，并以列表返回该列的所有数据。
    :param sheet: 工作表对象
    :param column_name: 列标题名称
    :return: 选择列的列表
    """

    try:
        # 获取列标题所在行
        header_row = sheet[1]

        # 根据列标题名称确定列索引
        column_index = None
        for cell in header_row:
            if cell.value == column_name:
                column_index = cell.column
                break

        # 如果未找到列标题，则抛出异常
        if column_index is None:
            raise ValueError(f"Column '{column_name}' not found in sheet '{sheet}'")

        # 遍历指定列的所有数据，从第二行开始
        column_list = []
        for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=column_index, max_col=column_index,
                                   values_only=True):
            if row[0] is not None:
                column_list.append(row[0])

        return column_list

    except Exception as e:
        print(f"An error occurred in get_column_list: {e}")
        return []  # 返回空列表表示发生了错误，无法获取列数据


def excel_get_column_str_custom(sheet, column_name, separator='、', end_mark='。'):
    """
    遍历指定工作簿中指定工作表的指定列，并返回该列的所有数据拼接成的字符串。
    :param sheet_name: 工作表对象
    :param column_name: 列标题名称
    :param separator: 单元格值之间的分隔符，默认为 '、'
    :param end_mark: 列字符串的结束标记，默认为 '。'
    :return: 拼接后的字符串
    """
    try:
        # 获取列标题所在行
        header_row = sheet[1]

        # 根据列标题名称确定列索引
        column_index = None
        for cell in header_row:
            if cell.value == column_name:
                column_index = cell.column
                break

        # 如果未找到列标题，则抛出异常
        if column_index is None:
            raise ValueError(f"Column '{column_name}' not found in sheet '{sheet_name}'")

        # 存储列数据的列表
        column_data = []

        # 遍历指定列的所有数据，从第二行开始
        for row in range(2, sheet.max_row + 1):
            cell_value = sheet.cell(row=row, column=column_index).value
            # 检查单元格的值是否为 None，如果是则跳过
            if cell_value is not None:
                column_data.append(str(cell_value))

        # 如果列表不为空，且最后一个元素不为空，则在最后一个元素后面添加结束标记
        if column_data and column_data[-1]:
            column_data[-1] += end_mark

        # 将列表中的数据拼接为一个字符串，用分隔符分隔
        column_string = f'{separator}'.join(column_data)

        return column_string

    except Exception as e:
        print(f"An error occurred in get_column_str_custom: {e}")
        return ""  # 返回空字符串表示发生了错误，无法获取列数据


def excel_get_columns_keyword_counts(sheet, platform_prefix='prefix', result_suffix='suffix'):
    """
    统计不同关键词的计数情况并汇总。
    :param sheet: Excel工作表对象
    :param platform_prefix: 平台列名称的前缀，默认为'prefix'
    :param result_suffix: 执行结果列名称的后缀，默认为'suffix'
    :return: 各个相同命名规则列的计数情况汇总列表
    """
    try:
        # 获取表头信息
        header_row = sheet[1]
        header = [cell.value for cell in header_row]

        # 初始化存储各个平台执行情况的列表
        list_platform_results = []

        # 遍历表头中所有列
        for i, col_name in enumerate(header):
            if col_name.startswith(platform_prefix) and col_name.endswith(result_suffix):
                try:
                    # 提取关键词名称
                    platform_name = col_name[len(platform_prefix):-len(result_suffix)]

                    # 初始化当前的统计情况字典
                    platform_results = {'platform': platform_name, 'results': {}}

                    # 找到当前列的索引
                    platform_column_index = i + 1

                    # 统计当前列的值及其出现次数
                    platform_counts = {}
                    for row in sheet.iter_rows(min_row=2, values_only=True):
                        result = row[platform_column_index - 1]  # 获取当前行的执行结果
                        if result not in platform_counts:
                            platform_counts[result] = 0
                        platform_counts[result] += 1

                    # 将统计结果保存到当前平台的执行情况字典中
                    platform_results['results'] = platform_counts

                    # 将当前平台的执行情况添加到结果列表中
                    list_platform_results.append(platform_results)

                except Exception as e:
                    print(
                        f"An error occurred while processing column '{col_name}': {e} occurred in get_columns_keyword_counts")

        return list_platform_results

    except Exception as e:
        print(f"An error occurred in get_columns_keyword_counts: {e} ")
        return []


def word_get_section_number(paragraph_text):
    """
    从段落文本中提取章节号
    """
    try:
        parts = paragraph_text.split()
        for part in parts:
            if part.replace('.', '').isdigit():
                return part
        return None
    except Exception as e:
        print(f"An error occurred in get_section_number: {e}")
        return None


def word_get_table_cell_text(table, row_index, column_index):
    """
    获取表格中指定单元格的文本值。
    :param table: 表格对象。
    :param row_index: 行索引（0开始）。
    :param column_index: 列索引（0开始）。
    :return: 单元格中的文本值。
    """
    try:
        return table.cell(row_index, column_index).text
    except AttributeError as e:
        print(f"AttributeError: {e} occurred in get_table_cell_text")
        return None
    except IndexError as e:
        print(f"IndexError: {e} occurred in get_table_cell_text")
        return None


def excel_get_sheets_name_row(source_excel):
    """
    获取excel的所有sheet的名称及其行数（默认去除标题行）
    :param source_excel: Excel 文件路径。
    :return: [{sheet_name}:{row_count}]
    """
    try:
        # 获取excel
        wb = openpyxl.load_workbook(source_excel)

        # 获取所有的sheet的名字
        sheet_names = wb.sheetnames

        # 初始化存储每个sheet数据行数的列表
        dict_sheet_name_rows = []

        # 遍历每个sheet
        for sheet_name in sheet_names:
            if sheet_name == 'Summary':  # Skip the Summary sheet
                continue
            sheet = wb[sheet_name]
            data_rows = 0
            # 统计非空行数
            for row in sheet.iter_rows(values_only=True):
                if not all(cell is None or cell == "" for cell in row):
                    data_rows += 1
            # 将工作表名称及其数据行数添加到字典中
            sheet_name_row = {sheet_name: data_rows - 1}  # 减去标题行
            dict_sheet_name_rows.append(sheet_name_row)

        return dict_sheet_name_rows

    except Exception as e:
        print(f"An error occurred in get_sheets_name_row: {e}")
        return None


def excel_get_sheet(source_excel, sheet_index):
    """
    获取工作簿中指定索引的工作表对象。
    :param source_excel:
    :param sheet_index: 表单索引。
    :return: 表单对象。
    """
    try:
        # 加载工作表
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有的工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引在有效范围内
        if 0 <= sheet_index < len(sheet_names):
            return workbook[sheet_names[sheet_index]]
        else:
            raise IndexError("Sheet index out of range.")
    except IndexError as e:
        print(f"An error occurred in get_sheet: {e}")
        return None


def excel_get_sheet_row_values(sheet, row_index):
    """
    获取指定工作表中某一行的数据。

    参数：
    sheet: openpyxl.Worksheet 对象，指定的工作表。
    row_index: int, 行索引，从 1 开始。

    返回值：
    list, 指定行的数据（所有单元格值）。
    """
    try:
        # 检查 sheet 是否为有效的工作表对象
        if sheet is None:
            raise ValueError("Invalid sheet object provided.")
        # 获取指定行的数据
        row_data = [cell.value for cell in sheet[row_index]]
        return row_data
    except IndexError:
        print("指定的行索引超出范围。")
        return []
    except Exception as e:
        print(f"获取行数据时发生错误: {e}")
        return []


def excel_get_sheet_col_values(sheet, col_index):
    """
    获取指定工作表中某一列的数据（兼容只读模式）。

    参数：
    sheet: openpyxl.Worksheet 对象，指定的工作表。
    col_index: int, 列索引，从 1 开始。

    返回值：
    list, 指定列的数据（所有单元格值）。
    """
    try:
        # 检查 sheet 是否为有效的工作表对象
        if sheet is None:
            raise ValueError("Invalid sheet object provided.")

        # 获取列数据（兼容只读模式）
        col_data = []
        for row in sheet.iter_rows(min_col=col_index, max_col=col_index, values_only=True):
            col_data.append(row[0])  # 每行的指定列值

        return col_data
    except IndexError:
        print("指定的列索引超出范围。")
        return []
    except Exception as e:
        print(f"获取列数据时发生错误: {e}")
        return []


def excel_get_sheet_name(sheet_data, sheet_index):
    """
    获取sheet的名称
    :param sheet_data: get_sheets_name_row()的返回值
    :param sheet_index: sheet的序号，以0开始
    :return:sheet名称
    """
    try:
        return [list(data.keys())[0] for data in sheet_data][sheet_index]
    except Exception as e:
        print(f"An error occurred in get_sheet_name: {e}")


def excel_get_sheet_row_col(sheet):
    """
    获取指定工作表的总行数和总列数。

    参数：
    sheet: openpyxl.Worksheet 对象，指定的工作表，excel_get_sheet()返回值。

    返回值：
    tuple, (总行数, 总列数)，如果出错则返回 (0, 0)。
    """
    try:
        # 检查 sheet 是否为有效的工作表对象
        if sheet is None:
            raise ValueError("Invalid sheet object provided.")
        # 获取行数和列数
        total_rows = sheet.max_row
        total_columns = sheet.max_column
        return total_rows, total_columns
    except Exception as e:
        print(f"获取行数和列数时发生错误: {e}")
        return 0, 0


def excel_insert_sheet_columns(source_excel, sheet_index, target_column_index, number_of_columns):
    """
    在指定的 Excel 工作表中，从目标列索引的后面插入指定数量的列。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 表单索引，从 0 开始。
    target_column_index: int, 目标列索引，从 1（A） 开始。
    number_of_columns: int, 需要插入的列数。

    返回值：
    bool, 插入操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 获取工作表对象
            sheet = workbook[sheet_names[sheet_index]]
            # 在目标列的下一列插入列
            insert_column_index = target_column_index + 1
            sheet.insert_cols(insert_column_index, number_of_columns)
            # 保存更改
            workbook.save(source_excel)
            print(
                f"已在工作表 '{sheet_names[sheet_index]}' 的第 {target_column_index} 列插入了 {number_of_columns} 列。")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"插入列时发生错误: {e}")
        return False


def excel_insert_sheet_rows(source_excel, sheet_index, target_row_index, number_of_rows):
    """
    在指定的 Excel 工作表中，从目标行索引的后面插入指定数量的行。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 表单索引，从 0 开始。
    target_row_index: int, 目标行索引，从 1 开始。
    number_of_rows: int, 需要插入的行数。

    返回值：
    bool, 插入操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 获取工作表对象
            sheet = workbook[sheet_names[sheet_index]]
            # 在目标行的下一行插入行
            insert_row_index = target_row_index + 1
            sheet.insert_rows(insert_row_index, number_of_rows)
            # 保存更改
            workbook.save(source_excel)
            print(f"已在工作表 '{sheet_names[sheet_index]}' 的第 {target_row_index} 行插入了 {number_of_rows} 行。")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"插入行时发生错误: {e}")
        return False


def excel_set_sheet_cell_value(source_excel, sheet_index, row, column, value):
    """
    在指定的 Excel 工作表的指定行列设置值。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 表单索引，从 0 开始。
    row: int, 目标单元格所在行，从 1 开始。
    column: int, 目标单元格所在列，从 1 开始。
    value: 任意, 要设置的值。

    返回值：
    bool, 设置操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 获取工作表对象
            sheet = workbook[sheet_names[sheet_index]]
            # 设置指定单元格的值
            sheet.cell(row=row, column=column, value=value)
            # 保存更改
            workbook.save(source_excel)
            print(f"已在工作表 '{sheet_names[sheet_index]}' 的第 {row} 行，第 {column} 列设置值为: {value}")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"设置值时发生错误: {e}")
        return False


def excel_delete_sheet_rows(source_excel, sheet_index, target_row_index, number_of_rows):
    """
    在指定的 Excel 工作表中，从目标行索引开始删除指定数量的行。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 表单索引，从 0 开始。
    target_row_index: int, 要删除的起始行索引，从 1 开始。
    number_of_rows: int, 要删除的行数。

    返回值：
    bool, 删除操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 获取工作表对象
            sheet = workbook[sheet_names[sheet_index]]
            # 删除指定行
            sheet.delete_rows(target_row_index+1, number_of_rows)
            # 保存更改
            workbook.save(source_excel)
            print(f"已在工作表 '{sheet_names[sheet_index]}' 的第 {target_row_index} 行开始删除 {number_of_rows} 行。")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"删除行时发生错误: {e}")
        return False

def excel_delete_sheet_columns(source_excel, sheet_index, target_column_index, number_of_columns):
    """
    在指定的 Excel 工作表中，从目标列索引开始删除指定数量的列。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 表单索引，从 0 开始。
    target_column_index: int, 要删除的起始列索引，从 1 开始。
    number_of_columns: int, 要删除的列数。

    返回值：
    bool, 删除操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 获取工作表对象
            sheet = workbook[sheet_names[sheet_index]]
            # 删除指定列
            sheet.delete_cols(target_column_index, number_of_columns)
            # 保存更改
            workbook.save(source_excel)
            print(
                f"已在工作表 '{sheet_names[sheet_index]}' 的第 {target_column_index} 列开始删除 {number_of_columns} 列。")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"删除列时发生错误: {e}")
        return False


def excel_add_sheet(source_excel, new_sheet_name):
    """
    在 Excel 文件中新增一个工作表。

    参数：
    source_excel: str, Excel 文件路径。
    new_sheet_name: str, 新工作表的名称。

    返回值：
    bool, 添加操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 检查是否已存在相同名称的工作表
        if new_sheet_name in workbook.sheetnames:
            print(f"工作表 '{new_sheet_name}' 已经存在。")
            return False
        # 添加新工作表
        workbook.create_sheet(title=new_sheet_name)
        # 保存更改
        workbook.save(source_excel)
        print(f"成功添加新工作表: '{new_sheet_name}'")
        return True
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"添加新工作表时发生错误: {e}")
        return False


def excel_delete_sheet_by_index(source_excel, sheet_index):
    """
    从 Excel 文件中删除指定索引的工作表。

    参数：
    source_excel: str, Excel 文件路径。
    sheet_index: int, 要删除的工作表索引，从 0 开始。

    返回值：
    bool, 删除操作是否成功。
    """
    try:
        # 加载工作簿
        workbook = openpyxl.load_workbook(source_excel)
        # 获取所有工作表名称
        sheet_names = workbook.sheetnames
        # 确保索引有效
        if 0 <= sheet_index < len(sheet_names):
            # 删除指定工作表
            sheet_name_to_delete = sheet_names[sheet_index]
            del workbook[sheet_name_to_delete]
            # 保存更改
            workbook.save(source_excel)
            print(f"成功删除工作表: '{sheet_name_to_delete}'")
            return True
        else:
            raise IndexError("Sheet index out of range.")
    except FileNotFoundError:
        print("Excel 文件未找到，请检查文件路径是否正确。")
        return False
    except Exception as e:
        print(f"删除工作表时发生错误: {e}")
        return False


def word_merge_cells(table, start_row, start_column, end_row, end_column):
    """
    合并表格中指定范围内的单元格。

    参数：
    table: 表格对象
    start_row: 起始行索引
    start_column: 起始列索引
    end_row: 结束行索引
    end_column: 结束列索引
    """
    # 检查起始行和列是否在有效范围内
    if start_row < 0 or start_row >= len(table.rows):
        raise ValueError("Start row index out of range occurred in merge_cells.")
    if start_column < 0 or start_column >= len(table.columns):
        raise ValueError("Start column index out of range occurred in merge_cells.")

    # 检查结束行和列是否在有效范围内
    if end_row < 0 or end_row >= len(table.rows):
        raise ValueError("End row index out of range occurred in merge_cells.")
    if end_column < 0 or end_column >= len(table.columns):
        raise ValueError("End column index out of range occurred in merge_cells.")

    # 合并单元格
    table.cell(start_row, start_column).merge(table.cell(end_row, end_column))


def word_merge_cells_by_column(table, column_index):
    """
    在表格中根据指定列合并相同值的行，并在合并后的单元格中保留唯一值
    :param table: 表格对象
    :param column_index: 要合并的列索引
    """
    try:
        prev_value = None
        merge_start = None

        for row_index, row in enumerate(table.rows):
            current_cell_value = row.cells[column_index].text.strip()

            if current_cell_value != prev_value:
                if merge_start is not None:
                    # Merge cells from merge_start to current row - 1 in the specified column
                    for i in range(merge_start, row_index):
                        # Keep only the value of the first row in the merged cells
                        table.cell(i, column_index).text = prev_value
                        table.cell(i, column_index).merge(table.cell(row_index - 1, column_index))

                # Update merge_start to the current row
                merge_start = row_index

            # Update prev_value for the next iteration
            prev_value = current_cell_value

        # Merge the last set of cells if needed
        if merge_start is not None and merge_start != len(table.rows):
            for i in range(merge_start, len(table.rows)):
                # Keep only the value of the first row in the merged cells
                table.cell(i, column_index).text = prev_value
                table.cell(i, column_index).merge(table.cell(len(table.rows) - 1, column_index))

    except Exception as e:
        print(f"An error occurred in merge_cells_by_column: {e}")


def word_set_number_to_roman(number):
    """
    将整数转换为罗马数字。
    """
    try:
        roman_numerals = {
            1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI', 7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X',
            20: 'XX', 30: 'XXX', 40: 'XL', 50: 'L', 60: 'LX', 70: 'LXX', 80: 'LXXX', 90: 'XC', 100: 'C',
            200: 'CC', 300: 'CCC', 400: 'CD', 500: 'D', 600: 'DC', 700: 'DCC', 800: 'DCCC', 900: 'CM', 1000: 'M'
        }
        if n in roman_numerals:
            return roman_numerals[n]
        else:
            return ''.join([roman_numerals[int(digit) * 10 ** i] for i, digit in enumerate(reversed(str(n)))][::-1])
    except Exception as e:
        print(f"An error occurred in set_number_to_roman: {e}")


def word_set_table_cell_text(table, row_index, column_index, text):
    """
    设置表格中指定单元格的文本内容。

    参数：
    table: 表格对象
    row_index: 行索引（0开始）
    column_index: 列索引（0开始）
    text: 要设置的文本内容
    """
    try:
        table.cell(row_index, column_index).text = text
    except Exception as e:
        print(f"An error occurred in set_table_cell_text: {e}")


def word_set_table_alignment(table, start_row=None, start_column=None, end_row=None, end_column=None,
                             horizontal_alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             vertical_alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """
    设置表格中指定范围内的单元格的文本对齐方式，默认情况下将所有单元格的文本设置为水平和垂直居中对齐。

    参数：
    table: 表格对象
    start_row: 起始行索引，默认为 None（默认第一行）
    start_column: 起始列索引，默认为 None（默认第一列）
    end_row: 结束行索引，默认为 None（默认最后一行）
    end_column: 结束列索引，默认为 None（默认最后一列）
    horizontal_alignment: 水平对齐方式，默认为水平居中对齐
    vertical_alignment: 垂直对齐方式，默认为垂直居中对齐
    """
    try:
        # 默认情况下，如果未提供范围，则将范围设置为整个表格
        if start_row is None:
            start_row = 0
        if start_column is None:
            start_column = 0
        if end_row is None:
            end_row = len(table.rows) - 1
        if end_column is None:
            end_column = len(table.columns) - 1

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if (start_row <= i <= end_row and start_column <= j <= end_column) or (
                        start_row > end_row and start_column > end_column):
                    cell.paragraphs[0].alignment = horizontal_alignment
                    cell.vertical_alignment = vertical_alignment

    except Exception as e:
        print(f"An error occurred in set_table_alignment: {e}")


def word_set_table_style(table, font_size=10.5, spacing_before=0):
    """
    设置表格的字体大小和段前间距。
    :param table: 表格对象。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    :return:
    """
    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(spacing_before)  # 设置段前间距
    except Exception as e:
        print(f"An error occurred in set_table_style: {e}")


def word_get_platform_results(list_test_objs, platform_prefix='prefix', result_suffix='suffix'):
    """
    测试获取各个测试项下，不同平台的执行结果情况
    :param list_test_objs:
    :return:
    """
    try:
        list_sheet_platform_results = []
        for sheet_obj in list_test_objs:
            sheet_name = sheet_obj.title
            sheet_platform_results = get_columns_keyword_counts(sheet_obj, platform_prefix=platform_prefix,
                                                                result_suffix=result_suffix)
            for item in sheet_platform_results:
                item['sheet_name'] = sheet_name  # 添加工作表名称到每个平台结果中
            list_sheet_platform_results.extend(sheet_platform_results)

        return list_sheet_platform_results
    except Exception as e:
        print(f"An error occurred in get_platform_results: {e} ")
