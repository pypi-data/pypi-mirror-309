"""Document file handler implementation"""

from pathlib import Path
from typing import Any, cast

import fitz  # PyMuPDF
import pandas as pd
from docx import Document

from ..base import FileHandler
from ..exceptions import FileError
from ..types import PDFDocument, PDFPage


class DocumentHandler(FileHandler):
    """Handler for document files"""

    SUPPORTED_FORMATS = {
        ".pdf": "PDF",
        ".docx": "DOCX",
        ".doc": "DOC",
        ".rtf": "RTF",
        ".odt": "ODT",
    }

    async def read_pdf(self, path: Path) -> dict[str, Any]:
        """Read PDF document"""
        try:
            doc = cast(PDFDocument, fitz.open(str(path)))
            content = {
                "pages": [],
                "metadata": doc.metadata,
                "toc": doc.get_toc(),
                "form_fields": doc.get_form_text_fields(),
            }

            for page in doc:
                pdf_page = cast(PDFPage, page)
                content["pages"].append(
                    {
                        "text": pdf_page.get_text(),
                        "images": self._extract_images(pdf_page),
                        "tables": self._extract_tables(pdf_page),
                        "links": pdf_page.get_links(),
                    },
                )

            return content
        except Exception as e:
            raise FileError(f"Failed to read PDF: {e!s}", cause=e)

    async def read_docx(self, path: Path) -> dict[str, Any]:
        """Read DOCX document"""
        try:
            doc = Document(str(path))
            content = {"paragraphs": [], "tables": [], "sections": [], "styles": []}

            # Extract paragraphs
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else "default"
                content["paragraphs"].append(
                    {
                        "text": para.text,
                        "style": style_name,
                        "runs": [
                            {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                            }
                            for run in para.runs
                        ],
                    },
                )

            # Extract tables
            for table in doc.tables:
                content["tables"].append([[cell.text for cell in row.cells] for row in table.rows])

            return content
        except Exception as e:
            raise FileError(f"Failed to read DOCX: {e!s}", cause=e)

    async def write_pdf(self, path: Path, content: dict[str, Any], **kwargs: Any) -> None:
        """Write PDF document"""
        try:
            doc = cast(PDFDocument, fitz.open())

            # Add pages
            for page_content in content.get("pages", []):
                page = doc.new_page()

                # Add text
                if "text" in page_content:
                    page.insert_text((72, 72), page_content["text"])

                # Add images
                for img in page_content.get("images", []):
                    rect = fitz.Rect(*img.get("rect", (72, 72, 300, 300)))
                    page.insert_image(rect, stream=img["data"])

            # Add metadata
            if "metadata" in content:
                doc.set_metadata(content["metadata"])

            # Save document
            doc.save(str(path), **kwargs)
            doc.close()

        except Exception as e:
            raise FileError(f"Failed to write PDF: {e!s}", cause=e)

    async def write_docx(self, path: Path, content: dict[str, Any], **kwargs: Any) -> None:
        """Write DOCX document"""
        try:
            doc = Document()

            # Add paragraphs
            for para in content.get("paragraphs", []):
                p = doc.add_paragraph()
                for run in para.get("runs", [{"text": para.get("text", "")}]):
                    r = p.add_run(run["text"])
                    r.bold = run.get("bold", False)
                    r.italic = run.get("italic", False)
                    r.underline = run.get("underline", False)

            # Add tables
            for table_data in content.get("tables", []):
                if table_data and table_data[0]:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = str(cell)

            doc.save(str(path))

        except Exception as e:
            raise FileError(f"Failed to write DOCX: {e!s}", cause=e)

    async def extract_text(
        self,
        path: Path,
        structured: bool = False,
    ) -> str | dict[str, Any]:
        """Extract text from document"""
        try:
            if path.suffix == ".pdf":
                doc = cast(PDFDocument, fitz.open(str(path)))
                if structured:
                    return {
                        "pages": [cast(PDFPage, page).get_text() for page in doc],
                        "metadata": doc.metadata,
                    }
                return " ".join(cast(PDFPage, page).get_text() for page in doc)

            if path.suffix == ".docx":
                doc = Document(str(path))
                if structured:
                    return {
                        "paragraphs": [p.text for p in doc.paragraphs],
                        "tables": [
                            [cell.text for cell in row.cells]
                            for table in doc.tables
                            for row in table.rows
                        ],
                    }
                return " ".join(p.text for p in doc.paragraphs)

            return "" if not structured else {"text": ""}

        except Exception as e:
            raise FileError(f"Failed to extract text: {e!s}", cause=e)

    async def extract_tables(
        self,
        path: Path,
        as_dataframe: bool = False,
    ) -> list[list[list[str]]] | list[pd.DataFrame]:
        """Extract tables from document"""
        try:
            tables = []

            if path.suffix == ".pdf":
                doc = cast(PDFDocument, fitz.open(str(path)))
                for page in doc:
                    pdf_page = cast(PDFPage, page)
                    tables.extend(self._extract_tables(pdf_page))

            elif path.suffix == ".docx":
                doc = Document(str(path))
                for table in doc.tables:
                    tables.append([[cell.text for cell in row.cells] for row in table.rows])

            if as_dataframe:
                return [pd.DataFrame(table) for table in tables]
            return tables

        except Exception as e:
            raise FileError(f"Failed to extract tables: {e!s}", cause=e)

    def _extract_images(self, page: PDFPage) -> list[dict[str, Any]]:
        """Extract images from PDF page"""
        images = []
        for img in page.get_images():
            xref = img[0]
            if page.parent:
                base = page.parent.extract_image(xref)
                if base:
                    images.append(
                        {
                            "data": base["image"],
                            "size": (base.get("width", 0), base.get("height", 0)),
                            "format": base.get("ext", ""),
                            "colorspace": base.get("colorspace", ""),
                        },
                    )
        return images

    def _extract_tables(self, page: PDFPage) -> list[list[list[str]]]:
        """Extract tables from PDF page"""
        tables = []
        # TODO: Implement table detection using page layout analysis
        return tables

    async def validate(self, path: Path) -> bool:
        """Validate document file"""
        try:
            if path.suffix.lower() not in self.SUPPORTED_FORMATS:
                return False

            if path.suffix == ".pdf":
                doc = fitz.open(str(path))
                doc.close()
            elif path.suffix == ".docx":
                Document(str(path))
            return True

        except Exception:
            return False
